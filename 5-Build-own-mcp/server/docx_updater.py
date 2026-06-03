from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy
import re


# ==============================
# PARSE CLEAN SKILL LINES
# ==============================

def parse_skill_lines(suggestions: str) -> list:
    """
    Parse the LLM output which is expected in clean format:
        Category Label: tool1, tool2, tool3

    Returns a list of (label, value) tuples.
    Skips any line that doesn't match this format (meta-text, blanks, etc.)
    """
    skills = []

    for line in suggestions.splitlines():
        line = line.strip()

        # Skip blank lines
        if not line:
            continue

        # Skip obvious meta/prose lines
        if line.lower().startswith(("add ", "note", "tip", "*add", "here", "below",
                                    "the following", "these ", "sprinkle")):
            continue

        # Strip leading bullet markers if LLM added them
        line = re.sub(r"^[-*•]\s*", "", line).strip()

        # Must contain a colon to be a valid skill line
        if ":" not in line:
            continue

        parts = line.split(":", 1)
        label = parts[0].strip()
        value = parts[1].strip()

        # Skip if label looks like a section header (all caps or too long)
        if not label or len(label) > 60:
            continue

        # Skip if value is empty or looks like a sentence (ends with ." or long prose)
        if not value:
            continue

        skills.append((label, value))

    return skills


# ==============================
# FIND TECHNICAL SKILLS BOUNDS
# ==============================

def find_technical_skills_bounds(doc: Document):
    """
    Returns (first_bullet_idx, last_bullet_idx) — the paragraph indices
    of the first and last bullets inside the Technical Skills section.
    """
    paragraphs = doc.paragraphs
    in_skills = False
    first_idx = None
    last_idx = None

    for i, para in enumerate(paragraphs):
        text = para.text.strip().upper()

        if "TECHNICAL SKILLS" in text:
            in_skills = True
            continue

        if in_skills:
            # Stop at next major section
            if text in ("PROJECTS", "CERTIFICATIONS", "ACHIEVEMENTS",
                        "EDUCATION", "WORK EXPERIENCE", "LANGUAGES", "REFERENCES"):
                break
            # Record non-empty paragraphs as skill bullets
            if para.text.strip():
                if first_idx is None:
                    first_idx = i
                last_idx = i

    return first_idx, last_idx


# ==============================
# BUILD A MATCHING BULLET
# ==============================

def make_skill_bullet(bold_label: str, value: str, reference_p) -> OxmlElement:
    """
    Clone the formatting from an existing bullet paragraph and replace its
    content with bold_label (bold) + value (normal weight).
    """
    new_p = copy.deepcopy(reference_p)

    # Remove all children except <w:pPr>
    for child in list(new_p):
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag != "pPr":
            new_p.remove(child)

    # Bold label run
    r_bold = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rPr.append(OxmlElement("w:b"))
    r_bold.append(rPr)
    t_bold = OxmlElement("w:t")
    t_bold.text = bold_label + ":"
    t_bold.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r_bold.append(t_bold)
    new_p.append(r_bold)

    # Normal value run
    r_val = OxmlElement("w:r")
    t_val = OxmlElement("w:t")
    t_val.text = " " + value
    t_val.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r_val.append(t_val)
    new_p.append(r_val)

    return new_p


# ==============================
# MAIN UPDATE FUNCTION
# ==============================

def update_cv_docx(cv_path: str, suggestions: str, output_path: str = None) -> str:
    """
    Replace the Technical Skills bullets in the CV with the clean
    skill lines returned by the LLM.

    Format expected from LLM:
        Machine Learning: TensorFlow, PyTorch, scikit-learn
        Deep Learning & LLMs: Fine-tuning, LLaMA 3, GPT-4
        ...

    Saves to output_path (defaults to <cv_stem>_updated.docx).
    """

    cv_path = Path(cv_path)
    if output_path is None:
        output_path = cv_path.parent / (cv_path.stem + "_updated.docx")
    else:
        output_path = Path(output_path)

    # ── Parse skill lines from LLM output ─────────────────────────────
    skill_lines = parse_skill_lines(suggestions)

    if not skill_lines:
        print("No valid skill lines found in LLM output — CV saved unchanged.")
        doc = Document(str(cv_path))
        doc.save(str(output_path))
        return str(output_path)

    print(f"\nReplacing Technical Skills section with {len(skill_lines)} clean bullets:")
    for label, value in skill_lines:
        print(f"  • {label}: {value[:60]}{'...' if len(value) > 60 else ''}")

    # ── Load document ──────────────────────────────────────────────────
    doc = Document(str(cv_path))

    # ── Find the existing skill bullets to replace ────────────────────
    first_idx, last_idx = find_technical_skills_bounds(doc)

    if first_idx is None:
        print("WARNING: Could not find Technical Skills section — CV saved unchanged.")
        doc.save(str(output_path))
        return str(output_path)

    # ── Get a reference bullet paragraph for formatting ───────────────
    reference_p = doc.paragraphs[first_idx]._p

    # ── Build all new bullet elements ─────────────────────────────────
    new_bullets = [
        make_skill_bullet(label, value, reference_p)
        for label, value in skill_lines
    ]

    # ── Get the XML body and locate old bullet paragraphs ────────────
    body = doc.element.body
    paragraphs = doc.paragraphs

    # Collect all old skill bullet <w:p> elements to remove
    old_ps = [paragraphs[i]._p for i in range(first_idx, last_idx + 1)]

    # Insert new bullets before the first old one, then remove old ones
    # Insert in reverse order using addprevious to preserve correct final order
    anchor = old_ps[0]
    for new_p in new_bullets:
        anchor.addprevious(new_p)

    for old_p in old_ps:
        old_p.getparent().remove(old_p)

    # ── Save ───────────────────────────────────────────────────────────
    doc.save(str(output_path))
    print(f"\nUpdated CV saved to: {output_path}")
    return str(output_path)